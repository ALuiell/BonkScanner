from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "reverse" / "Megabonk_Reverse_Quick_Tutorial_RU.docx"

ACCENT = RGBColor(0x1E, 0x5A, 0x8A)
TEXT = RGBColor(0x20, 0x25, 0x2B)
MUTED = RGBColor(0x5B, 0x65, 0x73)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_run(paragraph, text: str, *, bold: bool = False, color=None, size=None, font=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    if font is not None:
        run.font.name = font
        r_fonts = run._element.rPr.rFonts
        r_fonts.set(qn("w:eastAsia"), font)
    return run


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    add_run(p, text, color=TEXT, size=11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, "- ", bold=True, color=ACCENT, size=11)
        add_run(p, item, color=TEXT, size=11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, f"{index}. ", bold=True, color=ACCENT, size=11)
        add_run(p, item, color=TEXT, size=11)


def add_code(doc: Document, lines: list[str]) -> None:
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(2 if idx == 0 else 0)
        p.paragraph_format.space_after = Pt(0 if idx < len(lines) - 1 else 8)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, line, color=RGBColor(0x1B, 0x1F, 0x23), size=10.5, font="Consolas")


def add_callout(doc: Document, title: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    add_run(p1, title, bold=True, color=ACCENT, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_run(p2, text, color=TEXT, size=10.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].width = Inches(widths[i])
        set_cell_shading(hdr[i], "DCEAF5")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, bold=True, color=ACCENT, size=10.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            add_run(p, value, color=TEXT, size=10)
    doc.add_paragraph()


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT

    for name in ("Title", "Subtitle"):
        doc.styles[name].font.name = "Aptos Display"

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = ACCENT

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Aptos Display"
    h2.font.size = Pt(13.5)
    h2.font.bold = True
    h2.font.color.rgb = ACCENT


def build() -> Document:
    doc = Document()
    set_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    title.paragraph_format.space_after = Pt(16)
    add_run(title, "Краткий учебник по реверсу", bold=True, color=ACCENT, size=24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    add_run(
        subtitle,
        "Быстрый вход для новичка на примерах из проекта MegabonkReroll: память, указатели, цепочки, поиск значения и проверка результата.",
        color=MUTED,
        size=12,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, f"Версия документа: {date.today().isoformat()}", color=MUTED, size=10.5)

    doc.add_paragraph()
    add_callout(
        doc,
        "Что это за документ",
        "Это короткий учебник, чтобы быстро понять базовую механику reverse engineering в игровых процессах. Он не пытается покрыть все. Его цель: дать тебе рабочую картину и сразу привязать ее к реальным цепочкам из этого репозитория.",
        "F4F8FB",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_heading("1. Самая короткая суть реверса", level=1)
    add_paragraph(
        doc,
        "В нашем контексте реверс означает: найти, где игра хранит нужные данные, понять, как до них дойти, и сделать это так, чтобы путь работал не только сейчас, но и после новой сессии."
    )
    add_bullets(
        doc,
        [
            "Если нашел число в памяти, это только старт.",
            "Если понял, как добраться до него через стабильную цепочку, это уже решение.",
            "Если перенес эту цепочку в код и проверил на новой сессии, это production-ready вариант.",
        ],
    )

    doc.add_heading("2. Четыре базовых слова", level=1)
    add_table(
        doc,
        ["Слово", "Что значит", "Пример из проекта"],
        [
            ["Address", "конкретное место в памяти сейчас", "текущий live-адрес значения HP"],
            ["Pointer", "адрес, который хранит другой адрес", "root указывает на owner_stats"],
            ["Offset", "смещение внутри объекта", "+0x40, +0x18, +0x2C"],
            ["Pointer chain", "маршрут из нескольких указателей", "GameAssembly -> class_ptr -> static_fields -> root"],
        ],
        widths=[1.3, 2.6, 2.7],
    )
    add_callout(
        doc,
        "Запомнить проще так",
        "Address это точка. Pointer это стрелка на другую точку. Offset это шаг внутри объекта. Pointer chain это маршрут из стрелок и шагов.",
        "EEF6FB",
    )

    doc.add_heading("3. Почему нельзя верить одному адресу", level=1)
    add_paragraph(
        doc,
        "Потому что большинство красивых адресов из Cheat Engine живут только в текущей сессии. После перезапуска игры адрес изменится. Поэтому в коде проекта используются не такие одноразовые адреса, а module-relative точки входа и цепочки от них."
    )
    add_code(
        doc,
        [
            "Плохо:",
            "0x1A3F7C9E120   <- работает только сейчас",
            "",
            "Хорошо:",
            "GameAssembly.dll + 0x2F6A4B8",
            "-> class_ptr",
            "-> +0xB8 static fields",
            "-> +0x40 owner_stats",
        ],
    )

    doc.add_heading("4. Типовая схема памяти игры", level=1)
    add_code(
        doc,
        [
            "GameAssembly.dll + TYPE_INFO_OFFSET",
            "-> class_ptr",
            "-> static fields",
            "-> singleton / root object",
            "-> field",
            "-> container / dictionary / entries",
            "-> final value",
        ],
    )
    add_paragraph(
        doc,
        "Очень часто путь выглядит именно так. Сначала вход через модуль, потом через класс, потом через статические поля к текущему объекту, а уже дальше к таблице, словарю или конкретному float/int значению."
    )

    doc.add_heading("5. Пример 1: как у нас читаются player stats", level=1)
    add_paragraph(
        doc,
        "В `player_stats.py` путь к итоговым стартам игрока начинается от стабильной точки в `GameAssembly.dll`."
    )
    add_code(
        doc,
        [
            "GameAssembly.dll + 0x2F6A4B8",
            "-> class_ptr",
            "-> +0xB8 static fields",
            "-> +0x0 root",
            "-> +0x40 owner_stats / PlayerStatsNew",
            "-> +0x10 stats_context",
            "-> +0x18 entries",
            "-> +0x2C + stat_id * 0x10",
        ],
    )
    add_bullets(
        doc,
        [
            "До `entries` мы идем по объектам.",
            "На последнем шаге уже вычисляем позицию нужного stat slot.",
            "Это хороший пример: сначала найди контейнер, потом позицию внутри него.",
        ],
    )

    doc.add_heading("6. Пример 2: как у нас читается инвентарь", level=1)
    add_paragraph(
        doc,
        "Инвентарь в этом проекте сидит рядом с player stats, но требует уже понимания словаря."
    )
    add_code(
        doc,
        [
            "GameAssembly.dll + 0x2F6A4B8",
            "-> class_ptr",
            "-> +0xB8 static fields",
            "-> +0x40 PlayerStatsNew",
            "-> +0xA0 inventory container",
            "-> +0x50 passive item dictionary",
            "-> +0x18 entries",
            "-> +0x20 first entry",
            "-> entry + 0x10 value",
            "-> item_value + 0x0 class_meta",
            "-> class_meta + 0x10 item name",
            "-> item_value + 0x18 stack count",
        ],
    )
    add_paragraph(
        doc,
        "Тут видно важную вещь: иногда мало найти объект. Нужно еще понять layout словаря. Иначе ты дойдешь до контейнера, но не сможешь достать нормальные имена и количества."
    )

    doc.add_heading("7. Пример 3: map stats и interactables", level=1)
    add_code(
        doc,
        [
            "GameAssembly.dll + 0x2FB5E68",
            "-> class_ptr",
            "-> +0xB8 static fields",
            "-> interactables dictionary",
            "-> +0x18 entries",
            "-> +0x20 count",
            "-> key at +0x8, value at +0x10",
            "-> value +0x10 max",
            "-> value +0x14 current",
        ],
    )
    add_bullets(
        doc,
        [
            "Ключи читаются как строки.",
            "Строки вроде `Chests` или `Microwaves` потом мапятся на внутренние enum значения.",
            "Это уже не просто чтение числа, а разбор структуры данных.",
        ],
    )

    doc.add_heading("8. Как искать нужное значение шаг за шагом", level=1)
    add_numbered(
        doc,
        [
            "Возьми значение, которое легко проверить глазами в игре.",
            "Найди его live-адрес через Cheat Engine.",
            "Подтверди, что это реально оно: значение меняется вместе с игрой.",
            "Посмотри, какой объект или код к нему обращается.",
            "Поднимись вверх до pointer chain.",
            "Постарайся довести цепочку до module-relative старта.",
            "Перезапусти игру и проверь, что путь остался верным, даже если live-адрес сменился.",
            "Только после этого переноси path в код.",
        ],
    )

    doc.add_heading("9. Как понять, что ты идешь правильно", level=1)
    add_bullets(
        doc,
        [
            "Значение совпадает с UI.",
            "Оно обновляется во время игры.",
            "После новой сессии chain все еще работает.",
            "Каждый шаг цепочки можно объяснить словами: что это за объект и зачем этот offset.",
        ],
    )

    doc.add_heading("10. Самые частые ошибки новичка", level=1)
    add_table(
        doc,
        ["Ошибка", "Что на самом деле происходит"],
        [
            ["Нашел число и сразу хочу писать его в код", "это может быть одноразовый heap-адрес"],
            ["Значение похоже на правду, значит оно правильное", "это может быть stale pointer или base value"],
            ["Root найден, значит все готово", "возможно еще не понят layout словаря или таблицы"],
            ["После апдейта сломалось одно поле, значит сломалось все", "иногда root жив, а поехал только последний offset или stat id"],
        ],
        widths=[2.9, 3.1],
    )

    doc.add_heading("11. Когда чтения памяти мало и нужен hook", level=1)
    add_paragraph(
        doc,
        "Чтение памяти подходит для наблюдения: показать статы, считать предметы, увидеть состояние карты. Но когда нужно встроиться в поведение игры, попасть в точный момент выполнения или вызвать внутреннюю функцию, уже может понадобиться hook."
    )
    add_bullets(
        doc,
        [
            "В проекте пример такого уровня работы находится в `hook_loader.py` и `native/BonkHook`.",
            "Там важны не только offsets, но и сигнатуры байт, безопасный момент инжекта и готовность runtime.",
            "То есть hook обычно более хрупкий, чем обычное чтение памяти.",
        ],
    )

    doc.add_heading("12. Мини-шпаргалка", level=1)
    add_callout(
        doc,
        "Если запомнить только одно",
        "Не думай 'как найти адрес'. Думай 'как построить стабильный маршрут до данных'. Это лучший короткий способ понять reverse в этом проекте.",
        "FFF7E8",
    )
    add_bullets(
        doc,
        [
            "Сначала найди live-значение.",
            "Потом найди pointer chain.",
            "Потом проверь новую сессию.",
            "Потом перенеси в код.",
            "Если нужно управлять поведением игры, думай в сторону hook.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
