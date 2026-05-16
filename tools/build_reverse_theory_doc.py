from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "reverse" / "Megabonk_Reverse_Theory_RU.docx"

TITLE = "Реверс и память игры в MegabonkReroll"
SUBTITLE = (
    "Практическое мини-руководство по тому, с чем мы реально работали в проекте: "
    "адреса, pointer chains, offsets, стабильные root paths, live-значения, инвентарь, "
    "статы, hooks и типовой reverse workflow."
)

ACCENT = RGBColor(0x1E, 0x5A, 0x8A)
TEXT = RGBColor(0x20, 0x25, 0x2B)
MUTED = RGBColor(0x5B, 0x65, 0x73)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, color=None, size=None, font=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    if font is not None:
        run.font.name = font
        r_fonts = run._element.rPr.rFonts
        r_fonts.set(qn("w:eastAsia"), font)
    return run


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    add_run(p, text, color=TEXT, size=11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, "- ", bold=True, color=ACCENT, size=11)
        add_run(p, item, color=TEXT, size=11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, f"{index}. ", bold=True, color=ACCENT, size=11)
        add_run(p, item, color=TEXT, size=11)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(2 if idx == 0 else 0)
        p.paragraph_format.space_after = Pt(0 if idx < len(lines) - 1 else 8)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, line, color=RGBColor(0x1B, 0x1F, 0x23), size=10.5, font="Consolas")


def add_callout(doc: Document, title: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    add_run(p1, title, bold=True, color=ACCENT, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_run(p2, text, color=TEXT, size=10.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        if widths:
            hdr_cells[index].width = Inches(widths[index])
        set_cell_shading(hdr_cells[index], "DCEAF5")
        p = hdr_cells[index].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, bold=True, color=ACCENT, size=10.5)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if widths:
                cells[index].width = Inches(widths[index])
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            add_run(p, value, color=TEXT, size=10)

    doc.add_paragraph()


def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT

    for style_name in ("Title", "Subtitle"):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Aptos Display"
    heading1.font.size = Pt(19)
    heading1.font.bold = True
    heading1.font.color.rgb = ACCENT

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Aptos Display"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = ACCENT

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "Aptos"
    heading3.font.size = Pt(11.5)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0x2F, 0x45, 0x5C)


def add_cover_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(18)
    add_run(title, TITLE, bold=True, color=ACCENT, size=24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    subtitle.paragraph_format.left_indent = Inches(0.25)
    subtitle.paragraph_format.right_indent = Inches(0.25)
    add_run(subtitle, SUBTITLE, color=MUTED, size=12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    add_run(meta, "Основано на коде, документации и истории веток проекта MegabonkReroll.", color=TEXT, size=11)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta2, f"Собрано автоматически: {date.today().isoformat()}", color=MUTED, size=10.5)

    doc.add_paragraph()
    add_callout(
        doc,
        "Для чего этот текст",
        "Это не академический учебник по reverse engineering. Это практическое объяснение простым языком именно того набора приемов, который использовался в этом проекте: чтение памяти, поиск стабильных путей, отличия между временным адресом и рабочей цепочкой, а также когда нужно читать, а когда уже нужна нативная hook-логика.",
        "F4F8FB",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)


def build_document() -> Document:
    doc = Document()
    set_base_styles(doc)
    add_cover_page(doc)

    doc.add_heading("1. Что такое реверс в контексте этого проекта", level=1)
    add_body(
        doc,
        "В этом проекте под реверсом мы в основном имеем в виду не взлом ради взлома, а аккуратное понимание того, где игра хранит нужные данные и как до них безопасно добраться. Программа не рисует значения сама: она ищет процесс игры, подключается к памяти, находит GameAssembly.dll, а дальше по известным смещениям и указателям доходит до объектов со статистикой, состоянием карты, инвентарем и служебными контроллерами."
    )
    add_body(
        doc,
        "То есть наш практический реверс состоит из трех слоев. Первый слой: понять, какой именно объект в памяти отвечает за нужную сущность. Второй слой: превратить это в стабильную pointer chain, которая переживает перезапуск игры. Третий слой: оформить найденный путь в коде и убедиться, что он читает именно живые данные, а не случайный мусор или значение из старой сессии."
    )
    add_table(
        doc,
        ["Поверхность", "Где в проекте", "Зачем нужна", "Тип работы"],
        [
            ["Map stats / interactables", "game_data.py", "оценка карты, reroll logic", "чтение памяти"],
            ["Player stats", "player_stats.py", "live stats, snapshots, timeline", "чтение памяти"],
            ["Passive item inventory", "player_stats.py", "список предметов и стеков", "чтение памяти + разбор layout"],
            ["Native hook restart / toggles", "hook_loader.py + native/BonkHook", "рестарт, ready-state, live toggles", "инжект + export calls"],
        ],
        widths=[1.45, 1.5, 1.95, 1.25],
    )
    add_callout(
        doc,
        "Главная мысль",
        "Reverse здесь нужен затем, чтобы превратить игровые внутренности в стабильные данные для программы. Если найден только красивый адрес в Cheat Engine, работа еще не закончена. Цель не адрес сам по себе, а воспроизводимый путь до значения.",
        "EEF6FB",
    )

    doc.add_heading("2. Что мы увидели по истории веток и документации", level=1)
    add_body(
        doc,
        "История веток показывает, что проект развивался волнами. Отдельно шла линия live player stats и recordings, отдельно линия native hook и in-game toggles, отдельно рефакторинг GUI. Это полезно понимать: reverse здесь не был изолированной теорией, он постоянно кормил реальные фичи интерфейса и автоматизации."
    )
    add_bullets(
        doc,
        [
            "Ветка с live stats и recordings привязана к player stats path и инвентарю.",
            "Ветка с native hook settings связана уже не просто с чтением, а с управлением поведением игры через экспортированные функции DLL.",
            "Документы в docs/reverse фиксируют правило проекта: не хардкодить разовые live-адреса, а искать stable root path от GameAssembly.dll + offset.",
        ],
    )
    add_body(
        doc,
        "Именно поэтому почти все reverse-отчеты в проекте крутятся вокруг одной идеи: сначала находим рабочее live-значение, потом доказываем, что путь до него стабилен между сессиями, и только после этого переносим в production-код."
    )

    doc.add_heading("3. Базовые термины простым языком", level=1)
    add_table(
        doc,
        ["Термин", "Простое объяснение", "Как это выглядит у нас"],
        [
            ["Address", "конкретное место в памяти прямо сейчас", "например временный адрес значения HP в текущем запуске"],
            ["Pointer", "адрес, в котором лежит другой адрес", "root указывает на owner_stats, owner_stats указывает на stats_context"],
            ["Offset", "смещение от базового адреса", "+0x40, +0x18, +0x2C и так далее"],
            ["Pointer chain", "цепочка переходов от одного указателя к другому", "GameAssembly + offset -> class_ptr -> static_fields -> root -> object"],
            ["Module-relative address", "адрес не абсолютный, а относительно модуля", "GameAssembly.dll + 0x2F6A4B8"],
            ["Live address", "адрес, который красивый только в этой сессии", "после перезапуска меняется"],
            ["TypeInfo / static fields", "статическая точка входа в IL2CPP/класс", "class_ptr + 0xB8"],
            ["Hook", "вмешательство в выполнение функции", "BonkHook DLL вызывает exports и работает внутри процесса"],
        ],
        widths=[1.35, 2.95, 2.0],
    )
    add_callout(
        doc,
        "Почему люди часто путаются",
        "Потому что в Cheat Engine можно очень быстро найти одно нужное число и кажется, что задача решена. Но одно число еще не равно хорошему reverse. Хороший reverse начинается тогда, когда ты понимаешь, откуда оно берется, какой объект его хранит, и как заново прийти к нему в новой сессии.",
        "F7FAFD",
    )

    doc.add_heading("4. Как обычно устроены данные в памяти игры", level=1)
    add_body(
        doc,
        "Если говорить совсем грубо, картина часто выглядит так: есть модуль GameAssembly.dll, внутри него есть статические точки входа для игровых классов, у класса есть static fields, дальше через поля можно дойти до текущего singleton или active owner object, а уже внутри него лежат нужные нам таблицы, словари, списки или простые поля."
    )
    add_code_block(
        doc,
        [
            "GameAssembly.dll + TYPE_INFO_OFFSET",
            "-> class_ptr",
            "-> class_ptr + 0xB8  (static fields)",
            "-> root or singleton instance",
            "-> field A",
            "-> field B",
            "-> entries / dictionary / value",
        ],
    )
    add_body(
        doc,
        "Когда значение простое, на последнем шаге мы читаем int или float. Когда структура сложнее, на конце может быть словарь. Тогда мало найти root: нужно еще понять layout словаря, где count, где entries, с какого смещения начинается первый entry, где key, где value, и что лежит внутри value."
    )
    add_bullets(
        doc,
        [
            "Для простых чисел программа читает `read_i32`, `read_float`, `read_u8`.",
            "Для строк используется `read_mono_string` или `read_ascii_string` в зависимости от формата.",
            "Для переходов по цепочке используется `read_ptr`.",
        ],
    )

    doc.add_heading("5. Самая важная идея: адрес не равен пути", level=1)
    add_body(
        doc,
        "Это, пожалуй, главный принцип всего проекта. Если ты в одной сессии нашел адрес `0x1A3F...`, это еще не повод писать его в код. Такой адрес обычно живет только до перезапуска процесса. После новой загрузки карты, нового ранa или полного рестарта игры он почти наверняка изменится."
    )
    add_body(
        doc,
        "Поэтому хорошая цель звучит не так: найти адрес Max HP. Правильная цель звучит так: найти стабильный путь до объекта, в котором лежит final Max HP, начиная с module-relative точки вроде `GameAssembly.dll + 0x2F6A4B8`."
    )
    add_table(
        doc,
        ["Плохой результат", "Хороший результат"],
        [
            ["'Вот адрес, он сейчас меняется как надо'", "'Вот цепочка от GameAssembly.dll, и она переживает новую сессию'"],
            ["'Значение похоже на правду'", "'Значение совпало с UI, обновляется в игре и переживает fresh run'"],
            ["'Cheat Engine что-то нашел'", "'Понятно, какой это объект и почему смещения именно такие'"],
        ],
        widths=[3.35, 3.35],
    )

    doc.add_heading("6. Реальные pointer chains из этого проекта", level=1)
    doc.add_heading("6.1. Player stats", level=2)
    add_body(
        doc,
        "Одна из самых показательных цепочек в проекте находится в `player_stats.py`. Она ведет к таблице final/effective player stats, из которой UI берет Max HP, Armor, Luck, XP Gain и другие значения."
    )
    add_code_block(
        doc,
        [
            "GameAssembly.dll + 0x2F6A4B8",
            "-> class_ptr",
            "-> +0xB8   static fields",
            "-> +0x0    root",
            "-> +0x40   owner_stats / PlayerStatsNew",
            "-> +0x10   stats_context",
            "-> +0x18   entries",
            "-> +0x2C + stat_id * 0x10   value",
        ],
    )
    add_body(
        doc,
        "Здесь особенно хорошо видно, как chain делится на две части. Первая часть доводит нас до корневого объекта со статистикой. Вторая часть уже адресует конкретный stat slot через формулу `base + stat_id * slot_size`. Это типичная ситуация: сначала добираемся до контейнера, потом вычисляем позицию внутри него."
    )

    doc.add_heading("6.2. Passive item inventory", level=2)
    add_body(
        doc,
        "Инвентарь использует почти тот же root, что и player stats. Это удобный пример того, как один и тот же корневой объект кормит сразу несколько фич."
    )
    add_code_block(
        doc,
        [
            "GameAssembly.dll + 0x2F6A4B8",
            "-> class_ptr",
            "-> +0xB8 static fields",
            "-> +0x0  root",
            "-> +0x40 PlayerStatsNew",
            "-> +0xA0 inventory container",
            "-> +0x50 passive item dictionary",
            "-> +0x18 entries",
            "-> +0x20 first entry",
            "-> entry + 0x10 value",
            "-> item_value + 0x0 class_meta",
            "-> class_meta + 0x10 ASCII class name",
            "-> item_value + 0x18 stack count",
        ],
    )
    add_body(
        doc,
        "Тут уже видно, почему reverse редко заканчивается на одном указателе. Нужно было понять не только root, но и layout словаря: где count, где entries, как устроен entry, где внутри item лежит имя класса и где лежит количество стаков."
    )

    doc.add_heading("6.3. Map stats / interactables", level=2)
    add_body(
        doc,
        "В `game_data.py` используется отдельный type info path для словаря interactables. Дальше ключи читаются как Mono-строки, а значения как контейнеры с полями current и max."
    )
    add_code_block(
        doc,
        [
            "GameAssembly.dll + 0x2FB5E68",
            "-> class_ptr",
            "-> +0xB8 static fields",
            "-> interactables dictionary",
            "-> +0x18 entries",
            "-> +0x20 count",
            "-> each entry: key at +0x8, value at +0x10",
            "-> value +0x10 max, value +0x14 current",
        ],
    )
    add_body(
        doc,
        "Это еще один важный паттерн: сначала нужно уметь прочитать структуру контейнера, а потом правильно интерпретировать ключи. В нашем случае ключи вроде `Boss Curses`, `Chests`, `Microwaves` мапятся на внутренний enum `MapStat`."
    )

    doc.add_heading("6.4. Hook readiness и why hooks are different", level=2)
    add_body(
        doc,
        "В `hook_loader.py` уровень работы уже другой. Там недостаточно просто читать поля. Код проверяет, что `GameAssembly.dll` загружена, что нужные функции находятся в исполнимой памяти, что их байтовая сигнатура совпадает с ожидаемой, и что singleton вроде `AlwaysManager.Instance` уже инициализирован."
    )
    add_code_block(
        doc,
        [
            "GameAssembly.dll + 0x2F6BAA8",
            "-> AlwaysManager TypeInfo",
            "-> +0xB8 static fields",
            "-> +0x0 AlwaysManager.Instance",
        ],
    )
    add_body(
        doc,
        "Это хороший пример разницы между read path и hook path. Для чтения нам часто хватает 'объект существует и поле не ноль'. Для hook-инжекта требования строже: функция должна быть именно той функцией, сигнатура должна совпасть, память должна быть исполнимой, а игра должна быть в безопасном состоянии."
    )

    doc.add_heading("7. Как искать то, что нужно: типовой reverse workflow", level=1)
    add_body(
        doc,
        "Ниже самая практическая схема. Это именно тот шаблон мышления, который полезен при работе с такими проектами."
    )
    add_numbered(
        doc,
        [
            "Выбери наблюдаемое значение. Лучше всего то, что легко меняется и видно в UI: HP, Luck, количество предметов, opacity setting, счетчик сундуков.",
            "Найди live-значение в Cheat Engine или другом инструменте. Если это число, фильтруй по changed/unchanged, increased/decreased, exact value.",
            "Подтверди, что ты действительно попал в нужное место: меняй значение в игре, смотри реакцию UI, ищи who writes / who accesses this address.",
            "Пойми, что это за объект. Не просто адрес, а какой владелец этого поля, какой контейнер, какая таблица, словарь или class instance.",
            "Поднимись вверх до стабильного root path. Идеально, если начало будет module-relative: `GameAssembly.dll + offset`.",
            "Собери pointer chain по шагам и подпиши каждый шаг человеческим смыслом: class_ptr, static_fields, owner, entries, value.",
            "Проверь новую сессию. Перезапусти игру или начни новый run и убедись, что live-адреса изменились, но chain все еще приводит к правильному объекту.",
            "Только после этого переноси путь в код, добавляй защиту от нулевых указателей и обрабатывай ошибки чтения.",
            "Сверь результат с игровым UI и с поведением фичи. Для проекта это особенно важно: значение должно совпасть не только один раз, но и в живом рантайме, и в snapshots/recordings.",
        ],
    )
    add_callout(
        doc,
        "Золотое правило поиска",
        "Ищи не просто число. Ищи объяснимую модель: что это за объект, как ты в него пришел, почему эта цепочка переживет новую сессию, и что сломается первым после обновления игры.",
        "EEF7F1",
    )

    doc.add_heading("8. Pointermaps, chains и подтверждение между сессиями", level=1)
    add_body(
        doc,
        "В документах проекта прямо отмечено, что сырые pointer scans без pointermaps часто создают много шума. Для стабильного пути полезно сравнивать минимум две сессии: в первой сохранить pointermap, во второй снова найти то же значение и затем сравнить карты. Так отсекаются случайные одноразовые маршруты через heap."
    )
    add_bullets(
        doc,
        [
            "Если live-адрес поменялся, а твоя chain все еще дошла до правильного значения, это хороший знак.",
            "Если root остался тем же, но часть данных стала мусором, значит проблема может быть в layout контейнера, последнем offset или mapping stat ids.",
            "Если даже начало path перестало резолвиться, скорее всего игра обновилась и сдвинулись type info offsets или static entry points.",
        ],
    )

    doc.add_heading("9. Как все связано между собой в этом проекте", level=1)
    add_body(
        doc,
        "Ниже полезно смотреть на систему не как на набор магических смещений, а как на конвейер."
    )
    add_table(
        doc,
        ["Шаг", "Что происходит", "Где это видно"],
        [
            ["1", "Программа цепляется к процессу игры", "memory.py / ProcessMemory"],
            ["2", "Находит базу модуля GameAssembly.dll", "module_base_address / module_offset"],
            ["3", "Через type info и static fields выходит на нужный singleton/root", "game_data.py, player_stats.py, hook_loader.py"],
            ["4", "Читает поля, словари, entries, строки и числа", "read_ptr, read_i32, read_float, read_mono_string"],
            ["5", "Преобразует сырые данные в доменную модель программы", "MapStat, PlayerStatValue, snapshots, UI labels"],
            ["6", "Использует это в логике reroll, live stats или hook flow", "logic.py, gui.py, hook_loader.py"],
        ],
        widths=[0.55, 3.85, 2.3],
    )
    add_body(
        doc,
        "То есть reverse не живет отдельно. Он напрямую влияет на UX: если chain плохая, UI покажет `--`, snapshots запишут пустоту, reroll logic начнет принимать плохие решения, а hook может пытаться инжектиться слишком рано."
    )

    doc.add_heading("10. Как понять, что найдено не то", level=1)
    add_bullets(
        doc,
        [
            "Значение выглядит правдоподобно, но не меняется во время игры. Часто это stale pointer или base value вместо final runtime value.",
            "После новой сессии все развалилось. Значит путь был привязан к heap-адресу, а не к стабильному root.",
            "Часть статов читается, часть нет. Возможно, root правильный, но stat table layout или stat ids уже изменились.",
            "Counts нормальные, а имена предметов мусор. Обычно проблема в class metadata path или string format.",
            "Hook раньше работал, а теперь нет. Значит нужно перепроверять не только offsets, но и сигнатуру байт, readiness state и timing injection.",
        ],
    )

    doc.add_heading("11. Практический словарь мышления для тебя", level=1)
    add_body(
        doc,
        "Если упростить до очень бытового языка, полезно держать в голове такую модель."
    )
    add_bullets(
        doc,
        [
            "Модуль GameAssembly.dll: это как большой район, в котором лежат стартовые точки.",
            "TypeInfo offset: это как адрес здания класса.",
            "Static fields: это как стойка регистрации, через которую можно попасть к текущему общему экземпляру.",
            "Root / owner object: это уже нужный объект, который живет в текущем рантайме.",
            "Offset: это шаг внутри объекта, куда надо сдвинуться.",
            "Pointer chain: это маршрут по комнатам, а не одна дверь.",
            "Live address: это случайный номер стула в комнате сегодня. Завтра он будет другой.",
            "Stable path: это схема прохода по зданию, которая остается понятной после перезапуска.",
        ],
    )

    doc.add_heading("12. Когда хватит чтения памяти, а когда нужен hook", level=1)
    add_table(
        doc,
        ["Ситуация", "Обычно хватает", "Когда уже нужен hook"],
        [
            ["Нужно показать значения в UI", "чтение памяти", "редко"],
            ["Нужно записывать snapshots по таймеру", "чтение памяти", "если нужен ранний или событийный trigger"],
            ["Нужно узнать итоговые stat values", "чтение памяти по стабильной chain", "если значение не живет в удобной структуре"],
            ["Нужно управлять behavior игры", "нет", "обычно нужен hook или вызов метода внутри процесса"],
            ["Нужно надежно среагировать на момент события", "не всегда", "hook часто лучше, чем polling"],
        ],
        widths=[2.0, 2.15, 2.55],
    )
    add_body(
        doc,
        "Проще говоря: чтение памяти хорошо подходит для наблюдения. Hook нужен, когда нужно встроиться в выполнение игры, вызвать экспорт, переключить поведение или попасть в очень точный жизненный момент объекта."
    )

    doc.add_heading("13. Что читать в этом репозитории дальше", level=1)
    add_bullets(
        doc,
        [
            "`memory.py` — самый базовый слой: как читаются байты, указатели, строки и числа.",
            "`game_data.py` — лучший пример чтения карты и interactables через словарь.",
            "`player_stats.py` — лучший пример стабильной pointer chain до статов и пассивных предметов.",
            "`hook_loader.py` — лучший пример того, чем hook-логика отличается от обычного чтения.",
            "`docs/reverse/MEMORY_PATH_INDEX.md` — карта актуальных путей после обновлений.",
            "`docs/reverse/REVERSE_RECOVERY_GUIDE.md` — практический план действий, когда игра обновилась и все поехало.",
            "`docs/reverse/reports/*.md` — детальные reverse-отчеты по отдельным подсистемам.",
        ],
    )

    doc.add_heading("14. Короткое резюме", level=1)
    add_body(
        doc,
        "Если собрать все в одну фразу, то реверс в этом проекте — это поиск стабильного и объяснимого маршрута от `GameAssembly.dll + offset` до нужного игрового объекта, а затем аккуратное чтение или использование этого объекта в коде. Указатели, offsets и chains — не отдельная магия, а просто способ пройти от устойчивой точки входа к живому значению."
    )
    add_body(
        doc,
        "Самая частая ошибка новичка — спутать одноразовый live-адрес со стабильным решением. Самый полезный навык — уметь объяснить каждый шаг цепочки человеческими словами. Если ты можешь сказать: 'здесь мы входим в класс, здесь берем static fields, здесь идем к текущему owner, здесь попадаем в entries, здесь читаем float или item name', значит ты уже реально понимаешь reverse, а не просто повторяешь найденный адрес."
    )
    add_callout(
        doc,
        "Финальная опора",
        "Думай не 'как найти адрес', а 'как построить маршрут до данных и доказать, что он стабильный'. Это именно та логика, на которой держатся docs/reverse и текущая реализация проекта.",
        "FFF7E8",
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
